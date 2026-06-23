import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import io


st.set_page_config(layout="wide")
st.title("⚡ Energy Analysis Dashboard (Half-Hourly Data)")

# -----------------------------
# Upload file
# -----------------------------
files = st.file_uploader(
    "Upload energy data (single or multiple files)",
    type=["xlsx", "csv"],
    accept_multiple_files=True
)


if files:

    all_data = []

    for file in files:
        if file.name.endswith(".csv"):
            temp_df = pd.read_csv(file)
        else:
            temp_df = pd.read_excel(file, engine="openpyxl", header=None)
            
            # First row = intervals (1–48)
            headers = temp_df.iloc[0].tolist()
            
            # Create proper column names
            headers[0] = "date"   # first column is date
            
            temp_df.columns = headers
            
            # Remove header row from data
            temp_df = temp_df[1:]

            temp_df.columns = temp_df.columns.astype(str).str.strip()
            # ✅ Convert interval columns ONLY (exclude first column)
            new_cols = ["date"] + [str(i) for i in range(1, len(temp_df.columns))]
            temp_df.columns = new_cols
            temp_df = temp_df.rename(columns={temp_df.columns[0]: "date"})
            temp_df["date"] = pd.to_datetime(temp_df["date"], dayfirst=True, errors="coerce")
        
        # Detect format
        
        
        # ✅ Melt correctly
        temp_df = temp_df.melt(
            id_vars=["date"],
            var_name="interval",
            value_name="consumption"
        )
        
        # Convert interval (1–48)
        temp_df["interval"] = pd.to_numeric(temp_df["interval"], errors="coerce")
        
        # ✅ Build datetime properly
        temp_df["datetime"] = temp_df["date"] + pd.to_timedelta(
            (temp_df["interval"] - 1) * 30, unit="minutes"
        )
        
        # ✅ Clean consumption
        temp_df["consumption"] = pd.to_numeric(temp_df["consumption"], errors="coerce")
        
        temp_df = temp_df.dropna(subset=["datetime", "consumption"])



        
            # Convert date
        temp_df["date"] = pd.to_datetime(temp_df["date"], dayfirst=True, errors="coerce")
        
            # ✅ Convert interval column to numeric (1–48)
            
            # ✅ Clean interval values properly
        temp_df["interval"] = temp_df["interval"].astype(str).str.strip()
            
            # Extract number from string (handles '1', '1.0', '1 ')
        temp_df["interval"] = temp_df["interval"].str.extract(r'(\d+)')
            
        temp_df["interval"] = pd.to_numeric(temp_df["interval"], errors="coerce")

        
            # ✅ Convert interval → time (30-min slots)
        temp_df["datetime"] = temp_df["date"] + pd.to_timedelta(
            (temp_df["interval"] - 1) * 30, unit="minutes"
        )
        
        temp_df = temp_df.drop(columns=["date", "interval"])
        
        # Convert to numeric (CRITICAL FIX)
        temp_df["consumption"] = pd.to_numeric(temp_df["consumption"], errors="coerce")
            
        # Clean
        temp_df = temp_df.dropna(subset=["datetime", "consumption"])



        # ✅ ADD FUEL TYPE FROM FILE NAME
        if "gas" in file.name.lower():
            temp_df["fuel"] = "Gas"
        elif "elec" in file.name.lower() or "electric" in file.name.lower():
            temp_df["fuel"] = "Electricity"
        else:
            temp_df["fuel"] = "Unknown"

        all_data.append(temp_df)

    df = pd.concat(all_data)
    df = df.sort_values("datetime")

    fuel_options = df["fuel"].unique()
    
    selected_fuel = st.sidebar.selectbox(
        "Select Fuel Type",
        ["All"] + list(fuel_options)
    )
    
    if selected_fuel != "All":
        df = df[df["fuel"] == selected_fuel]


    meter_column = st.sidebar.selectbox(
        "Select Meter Column (if applicable)",
        ["None"] + list(df.columns)
    )

    if meter_column != "None":
        meters = df[meter_column].dropna().unique()
    
        selected_meter = st.sidebar.selectbox(
            "Select Meter",
            ["All"] + list(meters)
        )
    
        if selected_meter != "All":
            df = df[df[meter_column] == selected_meter]
        else:
            df = df.groupby("datetime")["consumption"].sum().reset_index()


    st.write("Data sample")
    st.write(df.head())
    
    st.write("Consumption stats")
    st.write(df["consumption"].describe())


    # -----------------------------
    # Standardise columns
    # -----------------------------
    
    df = df.dropna(subset=["datetime"])
    # Sort
    df = df.sort_values("datetime")
    
    # Feature engineering
    df["date"] = df["datetime"].dt.date
    df["hour"] = df["datetime"].dt.hour + df["datetime"].dt.minute / 60
    df["day"] = df["datetime"].dt.day_name()
    df["weekday"] = df["datetime"].dt.dayofweek
    df["is_weekend"] = df["weekday"] >= 5
    
    # -----------------------------
    # Key metrics
    # -----------------------------
    base_load = df["consumption"].quantile(0.1)
    avg_load = df["consumption"].mean()
    peak_load = df["consumption"].max()
        
    # Day vs night
    day_load = df[(df["hour"] >= 8) & (df["hour"] <= 18)]["consumption"].mean()
    night_load = df[(df["hour"] < 6)]["consumption"].mean()
        
    # Weekend comparison
    weekday_avg = df[df["is_weekend"] == False]["consumption"].mean()
    weekend_avg = df[df["is_weekend"] == True]["consumption"].mean()
        
    # Variability
    load_std = df["consumption"].std()
        
    report_items = []
    
    # -----------------------------
    # 1. Time Series
    # -----------------------------
    st.subheader("📈 Time Series")
    
    fig, ax = plt.subplots(figsize=(10,4))
    ax.plot(df["datetime"], df["consumption"])
    ax.set_ylabel("kWh")
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        The time series shows energy consumption ranging from approximately **{base_load:.1f} kWh to {peak_load:.1f} kWh**.  
        Average demand is **{avg_load:.1f} kWh**, indicating overall site usage.
            
        Daily cycling is clearly visible, suggesting structured operational hours.  
        {'Significant variation between peaks and troughs indicates strong operational influence on demand.' if peak_load > base_load * 2 else 'Relatively stable demand suggests more constant operational usage.'}
        """
    
        
    st.markdown(text)
        
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Time Series", text, img))
    
    
    # -----------------------------
    # 2. Average Daily Profile
    # -----------------------------
    st.subheader("📊 Average Daily Load Profile")
    
    avg_profile = df.groupby("hour")["consumption"].mean()
    
    fig, ax = plt.subplots()
    ax.plot(avg_profile.index, avg_profile.values)
    ax.set_xlabel("Hour of Day")
    ax.set_ylabel("Average kWh")
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        The average daily profile shows a **base load of ~{base_load:.1f} kWh** during low-use periods 
        and peak demand reaching **~{peak_load:.1f} kWh** during active hours.
            
        Daytime consumption averages **{day_load:.1f} kWh**, compared to night-time levels of **{night_load:.1f} kWh**.
            
        {'A strong increase during working hours indicates occupancy-driven demand.' if day_load > night_load * 1.5 else 'Limited variation suggests equipment may be running continuously.'}
        """
    
    st.markdown(text)
        
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Average Daily Load Profile", text, img))
    
    # -----------------------------
    # 3. Weekday vs Weekend
    # -----------------------------
    st.subheader("📅 Weekday vs Weekend")
    
    weekday = df[df["is_weekend"] == False].groupby("hour")["consumption"].mean()
    weekend = df[df["is_weekend"] == True].groupby("hour")["consumption"].mean()
    
    fig, ax = plt.subplots()
    ax.plot(weekday.index, weekday.values, label="Weekday")
    ax.plot(weekend.index, weekend.values, label="Weekend")
    ax.legend()
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        Average weekday consumption is **{weekday_avg:.1f} kWh**, compared to **{weekend_avg:.1f} kWh** on weekends.
        
        {'There is a clear reduction in weekend consumption, indicating reduced occupancy and operational activity.' if weekend_avg < weekday_avg * 0.8 else 'Weekend consumption remains relatively high, suggesting systems may be running unnecessarily outside working hours.'}
        """
    
    st.markdown(text)
    
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Weekday vs Weekend", text, img))
    
    # -----------------------------
    # 4. Heatmap
    # -----------------------------
    st.subheader("🔥 Heatmap")
    
    heatmap = df.pivot_table(
        index="date",
        columns="hour",
        values="consumption",
        aggfunc="mean"
    )
    
    fig, ax = plt.subplots(figsize=(12,6))
    sns.heatmap(heatmap, cmap="coolwarm", ax=ax)
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        The heatmap highlights consistent daily patterns, with higher consumption concentrated during peak hours.
        
        Load variability (standard deviation = **{load_std:.1f} kWh**) indicates 
        {'significant fluctuations in demand across the dataset.' if load_std > 10 else 'relatively stable consumption patterns.'}
        
        This visual is particularly useful for identifying anomalies or unusual spikes in demand.
        """
    
    st.markdown(text)
    
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Heatmap", text, img))
    
    # -----------------------------
    # 5. Load Duration Curve
    # -----------------------------
    st.subheader("⚡ Load Duration Curve")
    
    ldc = df["consumption"].sort_values(ascending=False).reset_index(drop=True)
    
    fig, ax = plt.subplots()
    ax.plot(ldc)
    ax.set_ylabel("kWh")
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        The load duration curve shows that peak demand reaches **{peak_load:.1f} kWh**, 
        while the base load remains around **{base_load:.1f} kWh**.
        
        {'The steep curve suggests short periods of high demand.' if peak_load > avg_load * 1.5 else 'The relatively flat curve suggests consistent energy use across the period.'}
        
        This indicates how frequently high loads occur and helps separate base load from operational demand.
        """
    
    st.markdown(text)
    
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Load Duration Curve", text, img))
        
    # -----------------------------
    # 6. Peak Demand
    # -----------------------------
    st.subheader("🔺 Daily Peak Demand")
    
    peak = df.groupby("date")["consumption"].max()
    
    fig, ax = plt.subplots()
    ax.plot(peak.index, peak.values)
    ax.set_ylabel("Peak kWh")
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        Daily peak demand reaches up to **{peak_load:.1f} kWh**, indicating the highest operational load on site.
        
        Monitoring peak demand is important for identifying unusually high consumption days and potential cost impacts.
        
        {'High peaks relative to average demand suggest opportunities to reduce maximum load.' if peak_load > avg_load * 1.5 else 'Peak demand is relatively stable compared to average usage.'}
        """
        
    st.markdown(text)
        
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Peak Demand", text, img))
    
    # -----------------------------
    # 7. Histogram
    # -----------------------------
    st.subheader("📉 Load Distribution")
    
    fig, ax = plt.subplots()
    ax.hist(df["consumption"], bins=40)
    ax.set_xlabel("kWh")
    st.pyplot(fig)
    
    text = f"""
        **Insight:**  
        The distribution shows most consumption values centred around **{avg_load:.1f} kWh**, 
        with a base load near **{base_load:.1f} kWh**.
        
        {'A wide spread of values indicates varied operational demand throughout the day.' if load_std > 10 else 'A narrow distribution suggests consistent energy usage.'}
        """
    
    st.markdown(text)
        
    # Save figure
    img = io.BytesIO()
    fig.savefig(img, format='png')
    img.seek(0)
        
    report_items.append(("Histogram", text, img))
    
    # -----------------------------
    # Recommendations Engine
    # -----------------------------
    recommendations = []
        
    if night_load > base_load * 1.2:
        recommendations.append(
            f"Reduce out-of-hours consumption. Night load ({night_load:.1f} kWh) is high relative to base load ({base_load:.1f} kWh)."
        )
        
    if weekend_avg > weekday_avg * 0.8:
        recommendations.append(
            "Investigate weekend usage. Consumption remains high outside normal working days."
        )
        
    if peak_load > avg_load * 1.5:
        recommendations.append(
            f"Peak demand is high ({peak_load:.1f} kWh). Consider load shifting or reducing peak usage."
        )
    
    # -----------------------------
    # Copilot Prompt Generator
    # -----------------------------
    copilot_prompt = f"""
    You are an energy analyst. Based on the following building energy data:
        
    - Base load: {base_load:.1f} kWh
    - Average load: {avg_load:.1f} kWh
    - Peak load: {peak_load:.1f} kWh
    - Daytime load: {day_load:.1f} kWh
    - Night load: {night_load:.1f} kWh
    - Weekday average: {weekday_avg:.1f} kWh
    - Weekend average: {weekend_avg:.1f} kWh
    - Load variability: {load_std:.1f}
        
    Write a professional energy report including:
    1. Key findings
    2. Inefficiencies
    3. Energy-saving recommendations
    4. Operational improvements
        
    Use formal report language.
    """
    
    st.subheader("🤖 Copilot Prompt (for advanced analysis)")
    st.code(copilot_prompt)
    
    
    # -----------------------------
    # Download charts-ready data
    # -----------------------------
        
    st.subheader("📄 Download Energy Report")
        
    def create_word_report():
        doc = Document()
        doc.add_heading("Energy Analysis Report", 0)
        
        for title, text, img in report_items:
            doc.add_heading(title, level=1)
            doc.add_paragraph(text)
        
            # Add image
            doc.add_picture(img, width=Inches(6))
        
        file_stream = io.BytesIO()
        doc.add_heading("Recommendations", level=1)
        for rec in recommendations:
            doc.add_paragraph(f"- {rec}")
            
            doc.add_heading("Copilot Prompts", level=1)
            
            doc.add_paragraph("Use the following prompts in Microsoft Copilot or other AI tools:")
            
            doc.add_paragraph(copilot_prompt)
            
            doc.add_paragraph("Additional prompts:")
            
            doc.add_paragraph(
                f"Identify energy savings from base load {base_load:.1f} kWh and peak {peak_load:.1f} kWh."
            )
            doc.save(file_stream)
            file_stream.seek(0)
            
        return file_stream
        
    word_file = create_word_report()
        
    st.download_button(
        label="Download Word Report",
        data=word_file,
        file_name="energy_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
